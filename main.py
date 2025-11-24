#!/usr/bin/env python3
"""
study-notes-summarizer (Hybrid mode) with logging, caching, OCR check, and parallel folder processing
"""

import os
import sys
import argparse
import logging
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
import PyPDF2
import glob
import io
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

# Optional: pdf2image for scanned PDFs
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
    PROJECT_POPPLER_BIN = Path(__file__).parent / "poppler" / "bin"
    PDF2IMAGE_POPPLER_PATH = str(PROJECT_POPPLER_BIN) if PROJECT_POPPLER_BIN.exists() else None
except Exception:
    PDF2IMAGE_AVAILABLE = False
    PDF2IMAGE_POPPLER_PATH = None

# Output formats
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load .env and OpenAI client
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ---------------------------
# Config
# ---------------------------
CHUNK_MAX_CHARS = 4500
FINAL_MODEL = "gpt-4.1"
TEXT_MODEL = "gpt-4o-mini"
IMAGE_MODEL = "gpt-4o"
STYLE_CHOICES = ["notion", "pastel", "cornell", "plain"]
MAX_WORKERS = 4  # threads for folder processing

# ---------------------------
# Logging
# ---------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

# ---------------------------
# Utilities
# ---------------------------
def chunk_text(text: str, max_chars=CHUNK_MAX_CHARS):
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + max_chars, length)
        if end < length:
            nl = text.rfind("\n", start, end)
            if nl > start + 50:
                end = nl
        chunks.append(text[start:end].strip())
        start = end
    return chunks

def safe_write_text(path: Path, content: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

def safe_write_bytes(path: Path, b: bytes):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "wb") as f:
        f.write(b)

# ---------------------------
# Extraction functions
# ---------------------------
def extract_text_from_pdf(pdf_path: Path) -> str:
    text_parts = []
    try:
        with open(pdf_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for i, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
                logger.debug(f"Extracted page {i}/{len(reader.pages)} chars={len(page_text)}")
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""
    return "\n\n".join(text_parts)

def convert_pdf_page_to_image_bytes(pdf_path: Path, page_number: int, dpi=150):
    if not PDF2IMAGE_AVAILABLE:
        return None
    try:
        images = convert_from_path(
            str(pdf_path),
            dpi=dpi,
            first_page=page_number,
            last_page=page_number,
            poppler_path=PDF2IMAGE_POPPLER_PATH
        )
        if not images:
            return None
        bio = io.BytesIO()
        images[0].save(bio, format="PNG")
        return bio.getvalue()
    except Exception as e:
        logger.warning(f"pdf2image failed for page {page_number}: {e}")
        return None

# ---------------------------
# OpenAI helpers
# ---------------------------
def summarize_with_text_model(text: str, system_prompt: str = None, model: str = TEXT_MODEL) -> str:
    system_prompt = system_prompt or "You create clear, exam-ready study notes with headings, bullets, definitions, and examples."
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text}
        ]
    )
    return resp.choices[0].message.content

def image_to_text_with_vision(image_bytes: bytes, model: str = IMAGE_MODEL) -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": "Extract all handwritten or printed text. Raw output only."},
                    {"type": "input_image", "image": image_bytes}
                ]
            }
        ]
    )
    return resp.choices[0].message.content

# ---------------------------
# Summarization pipeline
# ---------------------------
def chunk_summarize(text: str, chunk_max_chars=CHUNK_MAX_CHARS) -> str:
    chunks = chunk_text(text, max_chars=chunk_max_chars)
    summaries = []
    for i, chunk in enumerate(chunks):
        try:
            s = summarize_with_text_model(chunk)
        except Exception as e:
            logger.warning(f"Chunk {i+1} summarization failed, retrying: {e}")
            time.sleep(2)
            s = summarize_with_text_model(chunk)
        summaries.append(s)
    combined = "\n\n".join(summaries)
    try:
        final_prompt = f"Combine the following summaries into a cohesive study guide:\n\n{combined}"
        final_resp = client.chat.completions.create(
            model=FINAL_MODEL,
            messages=[
                {"role": "system", "content": "You produce polished, exam-ready study guides."},
                {"role": "user", "content": final_prompt}
            ]
        )
        return final_resp.choices[0].message.content
    except Exception as e:
        logger.warning(f"Final polish failed: {e}")
        return combined

# ---------------------------
# DOCX + Markdown
# ---------------------------
def save_summary_docx(summary: str, out_path: Path, style: str = "notion"):
    doc = Document()
    title = doc.add_paragraph("📚 Study Notes")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 55, 115)
    doc.add_paragraph()
    for line in summary.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("###"):
            h = doc.add_paragraph(line.replace("###", "").strip())
            h.style = "Heading 2"
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", "").strip(), style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.2
    doc.save(out_path)
    logger.info(f"DOCX saved: {out_path}")

def save_summary_markdown(summary: str, out_path: Path):
    safe_write_text(out_path, summary)
    logger.info(f"Markdown saved: {out_path}")

# ---------------------------
# File summarization with caching
# ---------------------------
def summarize_pdf_file(path: Path, style: str = "notion") -> str:
    out_path = path.parent / f"{path.stem}_summary.txt"
    if out_path.exists():
        logger.info(f"Cached summary found for {path.name}, skipping summarization.")
        return out_path.read_text()
    logger.info(f"Summarizing PDF: {path.name}")
    raw_text = extract_text_from_pdf(path)
    pages_text = raw_text.split("\n\n")
    need_images = [i+1 for i, t in enumerate(pages_text) if len(t.strip()) < 20]
    if need_images and PDF2IMAGE_AVAILABLE:
        logger.info(f"{len(need_images)} pages need OCR")
        for page_no in need_images:
            img_bytes = convert_pdf_page_to_image_bytes(path, page_no)
            if img_bytes:
                pages_text[page_no-1] = image_to_text_with_vision(img_bytes)
    raw_text = "\n\n".join(pages_text)
    master_summary = chunk_summarize(raw_text)
    save_summary_markdown(master_summary, path.parent / f"{path.stem}_summary.md")
    save_summary_docx(master_summary, path.parent / f"{path.stem}_summary.docx", style=style)
    safe_write_text(out_path, master_summary)
    return master_summary

def summarize_image_file(path: Path, style: str = "notion") -> str:
    out_path = path.parent / f"{path.stem}_summary.txt"
    if out_path.exists():
        logger.info(f"Cached summary found for {path.name}, skipping.")
        return out_path.read_text()
    logger.info(f"Summarizing Image: {path.name}")
    with open(path, "rb") as f:
        img_bytes = f.read()
    text = image_to_text_with_vision(img_bytes)
    master_summary = chunk_summarize(text)
    save_summary_markdown(master_summary, path.parent / f"{path.stem}_summary.md")
    save_summary_docx(master_summary, path.parent / f"{path.stem}_summary.docx", style=style)
    safe_write_text(out_path, master_summary)
    return master_summary

# ---------------------------
# Folder processing with parallelization
# ---------------------------
def summarize_folder(folder: Path, combine: bool = False, style: str = "notion"):
    files = sorted(
        Path(f) for f in glob.glob(str(folder / "**/*"), recursive=True)
        if Path(f).suffix.lower() in [".pdf", ".png", ".jpg", ".jpeg"]
    )
    if not files:
        logger.warning("No PDF/image files found.")
        return
    logger.info(f"Found {len(files)} files in {folder}")
    if combine:
        combined_raw = ""
        for f in files:
            if f.suffix.lower() == ".pdf":
                combined_raw += extract_text_from_pdf(f) + "\n\n"
            else:
                with open(f, "rb") as fh:
                    combined_raw += image_to_text_with_vision(fh.read()) + "\n\n"
        master_summary = chunk_summarize(combined_raw)
        save_summary_markdown(master_summary, folder / "combined_summary.md")
        save_summary_docx(master_summary, folder / "combined_summary.docx", style=style)
        safe_write_text(folder / "combined_summary.txt", master_summary)
        logger.info("Combined folder summary complete.")
        return
    # Parallel processing
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = []
        for f in files:
            if f.suffix.lower() == ".pdf":
                futures.append(executor.submit(summarize_pdf_file, f, style))
            else:
                futures.append(executor.submit(summarize_image_file, f, style))
        for future in as_completed(futures):
            try:
                future.result()
            except Exception as e:
                logger.warning(f"Error during parallel summarization: {e}")

# ---------------------------
# CLI
# ---------------------------
def build_arg_parser():
    p = argparse.ArgumentParser(description="Study Notes Summarizer (Hybrid mode)")
    p.add_argument("path", help="PDF file or folder")
    p.add_argument("--folder", action="store_true", help="Process folder")
    p.add_argument("--combine", action="store_true", help="Combine folder summaries")
    p.add_argument("--style", choices=STYLE_CHOICES, default="notion")
    return p

def main():
    parser = build_arg_parser()
    args = parser.parse_args()
    path = Path(args.path)
    if not path.exists():
        logger.error(f"Path does not exist: {path}")
        sys.exit(1)
    if args.folder:
        summarize_folder(path, combine=args.combine, style=args.style)
    else:
        if path.suffix.lower() == ".pdf":
            summarize_pdf_file(path, style=args.style)
        elif path.suffix.lower() in [".png", ".jpg", ".jpeg"]:
            summarize_image_file(path, style=args.style)
        else:
            logger.error("Unsupported file type. Use .pdf, .png, .jpg, or .jpeg")
            sys.exit(1)

if __name__ == "__main__":
    if not PDF2IMAGE_AVAILABLE:
        logger.warning("pdf2image not installed. Scanned PDFs will not be processed.")
    elif PDF2IMAGE_AVAILABLE and not PDF2IMAGE_POPPLER_PATH:
        logger.warning("Poppler not found. Scanned PDFs may fail. Place 'poppler/bin' in project folder.")
    main()
