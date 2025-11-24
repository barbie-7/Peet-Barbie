import os
import argparse
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
import PyPDF2
import glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# -------------------------------------------------
# Load API key
# -------------------------------------------------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# -------------------------------------------------
# Extract text from PDF
# -------------------------------------------------
def extract_text_from_pdf(pdf_path: Path) -> str:
    text = ""
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for i, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                print(f"📄 Extracted page {i}/{len(reader.pages)}")
    except Exception as e:
        print(f"❌ PDF extraction failed: {e}")
        exit(1)

    return text.strip()

# -------------------------------------------------
# Summarize text using high-quality exam-ready prompt
# -------------------------------------------------
def summarize_text(text: str) -> str:
    prompt = f"""
You are an expert university-level study-note creator.  
Transform the raw content below into **perfect exam-ready notes**.

Your output MUST follow this structure:
- **Main topic headings** (H2 style, e.g., "### Topic")
- Bullet points for every concept
- Clear definitions of key terms (**bold them**)
- Step-by-step explanations for processes
- Small examples to help understanding
- Highlight common mistakes, misconceptions, or exam traps
- Remove irrelevant/repeated content
- Make everything clear, structured, and optimized for memorization

RAW CONTENT:
{text}
"""

    try:
        print("🤖 Generating high-quality summary...")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"❌ Text summarization failed: {e}")
        exit(1)

# -------------------------------------------------
# Summarize image using GPT-4o Vision
# -------------------------------------------------
def summarize_image(image_path: Path) -> str:
    try:
        print(f"🖼️ Processing image: {image_path.name}")

        with open(image_path, "rb") as img:
            img_bytes = img.read()

        print("🤖 GPT-4o-Vision extracting content...")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": "Extract all handwritten or printed notes from this image. Keep accuracy extremely high."},
                        {"type": "input_image", "image": img_bytes}
                    ]
                }
            ]
        )

        extracted = response.choices[0].message.content
        return summarize_text(extracted)

    except Exception as e:
        print(f"❌ Image extraction failed: {e}")
        exit(1)

# -------------------------------------------------
# Save summary to DOCX with beautiful formatting
# -------------------------------------------------
def save_summary_as_docx(summary: str, output_file: Path):
    doc = Document()

    # Title
    title = doc.add_paragraph("📚 Study Notes Summary")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 50, 140)

    doc.add_paragraph()

    for line in summary.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("###"):
            p = doc.add_paragraph(line.replace("###", "").strip())
            p.style = "Heading 2"
        elif line.startswith("- "):
            p = doc.add_paragraph(line, style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.25

    doc.save(output_file)
    print(f"💾 Saved DOCX: {output_file}")

# -------------------------------------------------
# Summarize all files in folder
# -------------------------------------------------
def summarize_folder(folder_path: Path):
    files = [
        Path(f) for f in glob.glob(str(folder_path / "**/*"), recursive=True)
        if Path(f).suffix.lower() in [".pdf", ".png", ".jpg", ".jpeg"]
    ]

    if not files:
        print("❌ No files found.")
        return

    print(f"📁 Found {len(files)} files.\n")

    for f in files:
        print(f"📌 Processing: {f.name}")

        if f.suffix.lower() == ".pdf":
            text = extract_text_from_pdf(f)
            summary = summarize_text(text)
        else:
            summary = summarize_image(f)

        # Save text
        txt_file = f.parent / f"{f.stem}_summary.txt"
        with open(txt_file, "w", encoding="utf-8") as txt:
            txt.write(summary)

        # Save docx
        docx_file = f.parent / f"{f.stem}_summary.docx"
        save_summary_as_docx(summary, docx_file)

        print(f"✅ Completed {f.name}\n")

# -------------------------------------------------
# Combine all files into one big summary
# -------------------------------------------------
def summarize_folder_as_one(folder_path: Path):
    files = [
        Path(f) for f in glob.glob(str(folder_path / "**/*"), recursive=True)
        if Path(f).suffix.lower() in [".pdf", ".png", ".jpg", ".jpeg"]
    ]

    if not files:
        print("❌ No study materials found.")
        return

    print(f"📁 Combining {len(files)} files...")

    combined_raw = ""

    for f in files:
        print(f"📌 Adding: {f.name}")
        if f.suffix.lower() == ".pdf":
            combined_raw += extract_text_from_pdf(f) + "\n\n"
        else:
            combined_raw += summarize_image(f) + "\n\n"

    summary = summarize_text(combined_raw)

    txt_file = folder_path / "combined_summary.txt"
    docx_file = folder_path / "combined_summary.docx"

    with open(txt_file, "w", encoding="utf-8") as txt:
        txt.write(summary)

    save_summary_as_docx(summary, docx_file)

    print(f"🎉 Combined summary created!")

# -------------------------------------------------
# MAIN
# -------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Study Notes Summarizer")
    parser.add_argument("path", help="Path to file or folder")
    parser.add_argument("--folder", action="store_true")
    parser.add_argument("--combine", action="store_true")
    args = parser.parse_args()

    path = Path(args.path)
    if not path.exists():
        print("❌ Invalid path.")
        exit(1)

    if args.folder:
        if args.combine:
            summarize_folder_as_one(path)
        else:
            summarize_folder(path)
    else:
        if path.suffix.lower() == ".pdf":
            text = extract_text_from_pdf(path)
            summary = summarize_text(text)
        else:
            summary = summarize_image(path)

        # Save outputs
        txt_file = path.parent / f"{path.stem}_summary.txt"
        docx_file = path.parent / f"{path.stem}_summary.docx"

        with open(txt_file, "w", encoding="utf-8") as txt:
            txt.write(summary)

        save_summary_as_docx(summary, docx_file)

        print(summary)
